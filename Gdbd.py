import sys
from logging.handlers import TimedRotatingFileHandler
from pathlib import Path

from pony.orm import db_session, Database
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm
import time
import wx
import os
import logging as _syslog


def get_logger_factory(log_name='run'):
    """

    :param log_name:
    :return:
    """
    LOG_PATH = os.path.join(os.path.expanduser("~"), 'logs' + os.sep + "Gdbd")
    LOG_FILE_SUFFIX = '%Y-%m-%d'
    if not os.path.exists(LOG_PATH):
        os.mkdir(LOG_PATH)

    LOG_FMT = '[%(asctime)s] [%(process)d] [%(levelname)s] - %(module)s.%(funcName)s (%(filename)s:%(lineno)d) - ' \
              '%(message)s'

    formatter = _syslog.Formatter(LOG_FMT)

    logger = _syslog.getLogger()
    logger.setLevel(_syslog.INFO)

    file_handler = TimedRotatingFileHandler(
        filename=LOG_PATH + "/" + log_name + '.log', when='MIDNIGHT', interval=1
    )

    file_handler.suffix = LOG_FILE_SUFFIX

    file_handler.setFormatter(formatter)

    logger.addHandler(file_handler)

    stream_handler = _syslog.StreamHandler(sys.stdout)
    stream_handler.setFormatter(formatter)
    stream_handler.setLevel(_syslog.INFO)
    logger.addHandler(stream_handler)
    return logger


logger = get_logger_factory()


class MainFrame(wx.Frame):
    def __init__(self, p, t):
        wx.Frame.__init__(self, id=wx.ID_ANY, parent=p, title=t, size=(800, 500))
        # 数据库配置
        self.host_tc = None
        self.port_tc = None
        self.database_tc = None
        self.username_tc = None
        self.password_tc = None

        # 数据库连接
        self.db = None

        self.b = None
        self.font = wx.Font(16, wx.DEFAULT, wx.NORMAL, wx.NORMAL)
        self.InitUI()

    def InitUI(self):
        # 画一个面板
        panel = wx.Panel(self, -1)
        # 画上一个垂直块
        v_box = wx.BoxSizer(wx.VERTICAL)

        # 弹性网格布局 5 行 2列 边距是10
        fgs = wx.FlexGridSizer(5, 2, 20, 50)

        # 创建组件
        # # 创建第一列元素
        host = wx.StaticText(panel, -1, "连接地址")
        port = wx.StaticText(panel, -1, "端口号")
        database = wx.StaticText(panel, -1, "数据库")
        username = wx.StaticText(panel, -1, "用户名")
        password = wx.StaticText(panel, -1, "密码")
        self.b = wx.Button(panel, -1, "确认连接")

        # # 创建第二列元素
        self.host_tc = wx.TextCtrl(panel, -1, "localhost", style=wx.TE_PROCESS_ENTER, name="连接地址")
        self.port_tc = wx.TextCtrl(panel, -1, "3306", style=wx.TE_PROCESS_ENTER, name="端口号")
        self.database_tc = wx.TextCtrl(panel, -1, style=wx.TE_PROCESS_ENTER, name="数据库")
        self.username_tc = wx.TextCtrl(panel, -1, "root", style=wx.TE_PROCESS_ENTER, name="用户名")
        self.password_tc = wx.TextCtrl(panel, -1, style=wx.TE_PASSWORD | wx.TE_PROCESS_ENTER, name="密码")

        fgs.AddMany(
            [(host, 0, wx.EXPAND | wx.LEFT, 60),
             (self.host_tc, 0, wx.EXPAND | wx.RIGHT, 60),
             (port, 0, wx.EXPAND | wx.LEFT, 60),
             (self.port_tc, 0, wx.EXPAND | wx.RIGHT, 380),
             (database, 0, wx.EXPAND | wx.LEFT, 60),
             (self.database_tc, 0, wx.EXPAND | wx.RIGHT, 250),
             (username, 0, wx.EXPAND | wx.LEFT, 60),
             (self.username_tc, 0, wx.EXPAND | wx.RIGHT, 250),
             (password, 0, wx.EXPAND | wx.LEFT, 60),
             (self.password_tc, 0, wx.EXPAND | wx.RIGHT, 250)])

        v_box.Add(fgs, proportion=0, flag=wx.ALIGN_CENTER | wx.TOP | wx.BOTTOM, border=60)
        v_box.Add(self.b, 0, wx.CENTER | wx.ALL, 50)
        v_box.SetSizeHints(panel)
        panel.SetSizer(v_box)

        # 设置样式
        # 设置一个字体
        host.SetFont(self.font)
        port.SetFont(self.font)
        database.SetFont(self.font)
        username.SetFont(self.font)
        password.SetFont(self.font)

        # 绑定事件
        self.Bind(wx.EVT_BUTTON, self.ConnectDatabase, self.b)
        self.host_tc.Bind(wx.EVT_TEXT_ENTER, self.ConnectDatabase)
        self.port_tc.Bind(wx.EVT_TEXT_ENTER, self.ConnectDatabase)
        self.database_tc.Bind(wx.EVT_TEXT_ENTER, self.ConnectDatabase)
        self.username_tc.Bind(wx.EVT_TEXT_ENTER, self.ConnectDatabase)
        self.password_tc.Bind(wx.EVT_TEXT_ENTER, self.ConnectDatabase)

        # 展示
        self.Center()
        self.Show()

    def ConnectDatabase(self, event):
        """
        初始化数据库连接
        :param event:
        :return:
        """
        try:
            # 验证输入
            host = self.host_tc.Value
            port = self.port_tc.Value
            db_name = self.database_tc.Value
            user = self.username_tc.Value
            password = self.password_tc.Value
            tc_list = [self.host_tc, self.port_tc, self.database_tc, self.username_tc, self.password_tc]
            for t in tc_list:
                if not t.Value:
                    msg_dialog = wx.MessageDialog(self, f"选项 [{t.Name}] 不可为空", caption="校验提示",
                                                  style=wx.ICON_ERROR)
                    msg_dialog.ShowModal()
                    return

            logger.info("[数据库配置]连接地址:" + host)
            logger.info("[数据库配置]端口号:" + port)
            logger.info("[数据库配置]数据库名称:" + db_name)
            logger.info("[数据库配置]用户名:" + user)
            logger.info("[数据库配置]密码:" + password)
            self.db = Database()
            self.db.bind(provider='mysql', host=host, port=int(port), user=user,
                         passwd=password,
                         db=db_name)
            logger.info("[数据库配置]数据库连接初始化成功")
            ExportFrame(self, "选择表导出文档", self.db, db_name)
            self.Show(False)
            msg_dialog = wx.MessageDialog(self, "数据库连接成功", caption="连接数据库", style=wx.STAY_ON_TOP | wx.ICON_INFORMATION)
        except Exception as e:
            print(e.args)
            msg_dialog = wx.MessageDialog(self, "数据库连接失败:%s" % str(e.args), caption="连接数据库",
                                          style=wx.ICON_ERROR)
        msg_dialog.ShowModal()


class ExportFrame(wx.Frame):

    def __init__(self, parent, title, db, db_name):
        wx.Frame.__init__(self, parent, wx.ID_ANY, title, size=(1200, 800))
        # 数据库配置
        self.db_name = db_name
        # 数据库连接
        self.db = db
        # 数据库中所有的表
        self.db_tables = []
        # 已经选择的表
        self.selected_tables = None
        self.selected_table_list_str = None
        self.export_file_path = None
        # 列表框
        self.table_list_box = None
        # 字体
        self.font = wx.Font(16, wx.DEFAULT, wx.NORMAL, wx.NORMAL)
        self.InitUI()

    def InitUI(self):
        """
        初始主页面UI
        :return:
        """
        # 画一个水平块
        h_box = wx.BoxSizer(wx.HORIZONTAL)
        # 画两个垂直块
        v_box1 = wx.BoxSizer(wx.VERTICAL)
        v_box2 = wx.BoxSizer(wx.VERTICAL)

        label1 = wx.StaticText(self, id=-1, label="数据库中的所有表")
        search_Text = wx.TextCtrl(self, -1, style=wx.TE_PROCESS_ENTER)
        # 获取所有的表
        rows = self.GetTables()
        self.db_tables = [t[0] for t in rows]
        self.table_list_box = wx.ListBox(self, id=-1, style=wx.LB_EXTENDED,
                                         choices=self.db_tables)
        v_box1.Add(label1, 0, wx.EXPAND | wx.TOP, 10)
        v_box1.Add(search_Text, 0, wx.EXPAND | wx.BOTTOM | wx.TOP, 10)
        v_box1.Add(self.table_list_box, 1, wx.EXPAND | wx.BOTTOM, 10)

        # 右尺寸调整器
        label2 = wx.StaticText(self, -1, label="已选择的表")
        self.selected_tables = wx.TextCtrl(self, style=wx.TE_MULTILINE)
        label3 = wx.StaticText(self, -1, label="文件导出的位置")
        self.export_file_path = wx.TextCtrl(self, value=self.GetExportPath())
        export_button = wx.Button(self, -1, label="确认生成文档", size=(60, 20))

        v_box2.Add(label2, 0, wx.EXPAND | wx.TOP, 10)
        v_box2.Add(self.selected_tables, 1, wx.EXPAND | wx.TOP, 10)
        v_box2.Add(label3, 0, wx.EXPAND | wx.TOP, 30)
        v_box2.Add(self.export_file_path, 0, wx.EXPAND | wx.TOP, 10)
        v_box2.Add(export_button, 0, wx.SHAPED | wx.ALIGN_CENTER | wx.TOP | wx.BOTTOM, 15)

        h_box.Add(v_box1, 1, wx.EXPAND | wx.LEFT, 10)
        h_box.Add(v_box2, 3, wx.EXPAND | wx.LEFT | wx.RIGHT, 10)

        label1.SetFont(self.font)
        label2.SetFont(self.font)
        label3.SetFont(self.font)

        self.SetSizer(h_box, deleteOld=True)

        # 绑定事件
        self.Bind(wx.EVT_LISTBOX_DCLICK, self.SelectedTable, self.table_list_box)
        self.Bind(wx.EVT_TEXT_ENTER, self.OnCharChanged, search_Text)
        self.Bind(wx.EVT_BUTTON, self.ExportFile, export_button)
        self.Bind(wx.EVT_CLOSE, self._OnClose)

        self.Center()
        self.Show()

    @db_session
    def GetTables(self, search_content=None, include=None):
        """
        获取所有表格
        :return:
        """
        sql = "select table_name,table_comment from information_schema.tables where table_schema='%s' and " \
              "table_type='%s' " % (self.db_name, "BASE TABLE")
        if search_content:
            sql += " and table_name like '%s' " % ("%" + search_content + "%")
        if include:
            sql += " and table_name in (%s) " % ",".join(["'" + str(t) + "'" for t in include])
        logger.info(sql)
        res = self.db.execute(sql)
        rows = res.fetchall()
        if rows:
            return rows
        return []

    def SelectedTable(self, event):
        """
        获取列表选框已经选择的项
        :param event:
        :return:
        """
        selections = (event.GetEventObject().GetSelections())
        old_selected_table_list_str = self.selected_table_list_str.replace(" ",
                                                                           "") if self.selected_table_list_str else None
        old_selected_list = old_selected_table_list_str.split(",") if old_selected_table_list_str else []
        cur_selected_list = [self.db_tables[i] for i in selections if self.db_tables[i] not in old_selected_list]
        if not self.selected_table_list_str: self.selected_table_list_str = ""
        self.selected_table_list_str += ",  " + ",  ".join(
            cur_selected_list) if old_selected_list and cur_selected_list else ",  ".join(
            cur_selected_list)
        self.selected_tables.SetValue(self.selected_table_list_str)

    def GetExportPath(self):
        """
        获取文件导出位置
        :return:
        """
        filename = "数据字典%s.docx" % str(time.time()).split(".")[0]
        self.export_file_path = os.path.join(os.path.expanduser("~"), 'Downloads' + os.sep + filename)
        return self.export_file_path

    def ExportFile(self, event):
        """
        导出数据库字典文档
        :return:
        """
        try:
            selected_table_str = self.selected_tables.Value
            if not selected_table_str:
                msg_dialog = wx.MessageDialog(self, "请选择表", caption="导出提示",
                                              style=wx.ICON_ERROR)
                msg_dialog.ShowModal()
                return
            file = Path(self.export_file_path.GetValue())
            if file.exists():
                msg_dialog = wx.MessageDialog(self, f"该文件({self.export_file_path.GetValue()})已经存在,是否覆盖文件?",
                                              caption="文件已存在",
                                              style=wx.YES_NO | wx.STAY_ON_TOP | wx.ICON_ERROR)
                if msg_dialog.ShowModal() == wx.ID_NO:
                    return
            table_numbers = 1
            db_name = self.db_name
            # 所有已经选择的表
            table_name_list = selected_table_str.replace(" ", "").split(
                ",") if selected_table_str else None
            tables = self.GetTables(include=table_name_list)

            # 创建文件对象
            doc = Document(docx="./default.docx")
            doc.styles['Normal'].font.name = '宋体'
            # 编写表汇总
            doc.add_paragraph("1.1. 表汇总")
            # 添加一个表
            table = doc.add_table(rows=1, cols=2, style="Light Grid Accent 1")
            # 表格的定位居中
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            # 拿到第一行、并添加第一行的数据
            row = table.rows[0]
            row.cells[0].text = "表名"
            row.cells[1].text = "功能说明"
            if tables:
                for t in tables:
                    table_name, table_comment = t
                    rowa = table.add_row()
                    rowa.cells[0].text = table_name
                    rowa.cells[1].text = table_comment

            doc.add_paragraph("\n")
            doc.add_paragraph("1.2. 表")
            # 编写表
            if tables:
                for t in tables:
                    table_name, table_comment = t
                    # 添加表头
                    filed_table = self.AddFiledTable(doc, table_name, table_numbers)
                    # 添加字段
                    table_columns = self.GetTableColumns(table_schema=db_name, table_name=table_name,
                                                         exclude=[])
                    if table_columns:
                        n = 1
                        pk_cols = ""
                        in_cols = ""
                        for col in table_columns:
                            COLUMN_NAME, COLUMN_TYPE, IS_NULLABLE, COLUMN_KEY, COLUMN_COMMENT = col
                            frow = filed_table.add_row()
                            frow.cells[0].text = str(n)
                            frow.cells[1].text = COLUMN_NAME
                            frow.cells[2].text = COLUMN_TYPE
                            frow.cells[3].text = 'Y' if IS_NULLABLE == 'yes' else 'N'
                            frow.cells[4].text = 'Y' if COLUMN_KEY else 'N'
                            frow.cells[5].text = COLUMN_COMMENT
                            if COLUMN_KEY:
                                if COLUMN_KEY == 'PRI':
                                    # 主键字段
                                    pk_cols += COLUMN_NAME + ","
                                # 添加索引字段
                                in_cols += COLUMN_NAME + ","
                            n += 1
                        filed_table.cell(1, 1).text = pk_cols[0:len(pk_cols) - 1] if pk_cols.endswith(",") else pk_cols
                        filed_table.cell(3, 1).text = in_cols[0:len(in_cols) - 1] if in_cols.endswith(",") else in_cols
                    table_numbers += 1
                    doc.add_paragraph("\n")

            # 保存文档
            doc.save(self.export_file_path.GetValue())
            msg_dialog = wx.MessageDialog(self, "导出成功,文件导出位置:%s" % self.export_file_path.GetValue(), caption="导出成功",
                                          style=wx.STAY_ON_TOP | wx.ICON_INFORMATION)
            msg_dialog.ShowModal()
        except Exception as e:
            msg_dialog = wx.MessageDialog(None, "导出异常:%s" % str(e.args), caption="导出失败",
                                          style=wx.ICON_ERROR)
            msg_dialog.ShowModal()

    @db_session
    def GetTableColumns(self, table_schema, table_name, exclude=None):
        """
        获取指定表的字段信息
        :param table_schema: 数据库名称
        :param table_name: 表名称
        :param exclude 排除表字段
        :return:
        """
        sql = "select COLUMN_NAME,COLUMN_TYPE,IS_NULLABLE,COLUMN_KEY,COLUMN_COMMENT from information_schema.columns where table_schema='%s' and table_name='%s' " % (
            table_schema, table_name)
        if exclude:
            sql += " and COLUMN_NAME not in (%s) " % ",".join(["'" + str(t) + "'" for t in exclude])
        logger.info(sql)
        res = self.db.execute(sql)
        return res.fetchall()

    def AddFiledTable(self, doc, table_name, table_numbers):
        """
        添加字段描述表
        :param doc: document对象
        :param table_name:表名称
        :return:
        """
        doc.add_paragraph("（%s）" % str(table_numbers))
        filed_table = doc.add_table(rows=4, cols=6, style="Light Grid Accent 1")
        # 表格的定位 水平居中
        filed_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        filed_table.columns[0].width = Cm(1.5)
        filed_table.columns[1].width = Cm(2.5)
        filed_table.columns[3].width = Cm(1.5)
        filed_table.columns[4].width = Cm(1.5)
        rows = filed_table.rows
        for r in rows:
            cells = r.cells
            cells[1].merge(cells[5])
        rows = filed_table.rows
        rows[0].cells[0].text = "表名"
        rows[0].cells[1].text = table_name

        rows[1].cells[0].text = "主键"
        rows[1].cells[1].text = ""

        rows[2].cells[0].text = "其他排序字段"
        rows[2].cells[1].text = ""

        rows[3].cells[0].text = "索引字段"
        rows[3].cells[1].text = ""

        # 添加一行
        h_row = filed_table.add_row()
        h_row.cells[0].text = "序号"
        h_row.cells[1].text = "字段名称"
        h_row.cells[2].text = "数据类型（精度范围）"
        h_row.cells[3].text = "允许为空Y/N"
        h_row.cells[4].text = "唯一Y/N"
        h_row.cells[5].text = "约束条件/说明"
        return filed_table

    def _OnClose(self, event):
        self.Show(False)
        self.Parent.Show(True)
        event.GetEventObject().Show(False)

    def OnCharChanged(self, event):
        search_text = event.GetEventObject().Value
        rows = self.GetTables(search_content=search_text)
        _db_tables = [t[0] for t in rows]
        self.db_tables = _db_tables
        self.table_list_box.Clear()
        self.table_list_box.Append(_db_tables)
        event.Skip()


class TextEmptyValidator(wx.Validator):
    def __init__(self):
        wx.Validator.__init__(self)

    def Clone(self):
        return TextEmptyValidator()

    def Validate(self, win):  # 1 使用验证器方法
        return True

    def TransferToWindow(self):
        return True

    def TransferFromWindow(self):
        return True


class SearchTextValidator(wx.Validator):
    def __init__(self, all_db_table_list=None, export_frame=None):
        wx.Validator.__init__(self)
        print("111")
        self.all_db_table_list = all_db_table_list if all_db_table_list else []
        self.export_frame = export_frame
        self.Bind(wx.EVT_CHAR, self.OnCharChanged)  # 绑定字符输入事件
        self.search_text = ""

    def OnCharChanged(self, event):
        # 得到输入字符的 ASCII 码
        keycode = event.GetKeyCode()
        print(keycode)
        # 把 ASII 码 转成字符
        InputChar = chr(keycode)
        if InputChar:
            self.search_text += InputChar
            # 第一个字符为 .,非法，拦截该事件，不会成功输入
            _db_tables = [t for t in self.all_db_table_list if self.search_text in self.all_db_table_list]
            self.export_frame.db_tables = _db_tables
            event.Skip()
            return True
        return False

    def Clone(self):
        return TextEmptyValidator()

    def Validate(self, win):  # 1 使用验证器方法
        return True

    def TransferToWindow(self):
        return True

    def TransferFromWindow(self):
        return True


if __name__ == '__main__':
    try:
        app = wx.App()
        f = MainFrame(None, "生成数据字典程序")
        app.MainLoop()
    except Exception as e:
        logger.error(e.args, e)
        msg_dialog = wx.MessageDialog(None, "运行异常:%s" % str(e.args), caption="异常信息",
                                      style=wx.ICON_ERROR)
        msg_dialog.ShowModal()
